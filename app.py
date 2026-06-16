from flask import Flask, request, jsonify
from flask_cors import CORS
from kontur import kontur_bp
import io
import os

app = Flask(__name__)
CORS(app)
app.register_blueprint(kontur_bp)

def extract_text_from_pdf(file_bytes):
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        if text.strip():
            return text
        return extract_text_ocr(file_bytes)
    except Exception as e:
        return f"Ошибка PDF: {str(e)}"

def extract_text_ocr(file_bytes):
    try:
        import pytesseract
        from pdf2image import convert_from_bytes
        images = convert_from_bytes(file_bytes)
        text = ""
        for image in images:
            text += pytesseract.image_to_string(image, lang='rus+eng') + "\n"
        return text
    except Exception as e:
        return f"Ошибка OCR: {str(e)}"

def extract_text_from_docx(file_bytes):
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        text = ""
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " | "
                text += "\n"
        return text
    except Exception as e:
        return f"Ошибка DOCX: {str(e)}"

def extract_text_from_xlsx(file_bytes):
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
        text = ""
        for sheet in wb.worksheets:
            text += f"Лист: {sheet.title}\n"
            for row in sheet.iter_rows():
                row_text = []
                for cell in row:
                    if cell.value is not None:
                        row_text.append(str(cell.value))
                if row_text:
                    text += " | ".join(row_text) + "\n"
        return text
    except Exception as e:
        return f"Ошибка XLSX: {str(e)}"

@app.route('/parse', methods=['POST'])
def parse_document():
    if 'file' not in request.files:
        return jsonify({'error': 'Файл не найден'}), 400

    file = request.files['file']
    filename = file.filename.lower()
    file_bytes = file.read()

    if filename.endswith('.pdf'):
        text = extract_text_from_pdf(file_bytes)
    elif filename.endswith('.docx'):
        text = extract_text_from_docx(file_bytes)
    elif filename.endswith('.xlsx') or filename.endswith('.xls'):
        text = extract_text_from_xlsx(file_bytes)
    else:
        text = "Неподдерживаемый формат файла"

    if len(text) > 5000:
        text = text[:5000] + "...[текст обрезан]"

    return jsonify({'text': text, 'filename': file.filename})

@app.route('/health', methods=['GET'])
def health():
    return jsonify({'status': 'ok'})

@app.route('/zakupki', methods=['GET'])
def get_zakupki():
    try:
        import ftplib
        import zipfile
        import xml.etree.ElementTree as ET

        keyword = request.args.get('keyword', 'чиллер').lower()
        region = request.args.get('region', 'Moskva')
        results = []

        ftp = ftplib.FTP('ftp.zakupki.gov.ru', timeout=30)
        ftp.login('free', 'free')
        ftp.cwd(f'/out/published/{region}/notifications/')

        files = ftp.nlst()
        zip_files = sorted([f for f in files if f.endswith('.zip')])

        if not zip_files:
            return jsonify({'error': 'Файлы не найдены', 'region': region}), 404

        last_file = zip_files[-1]

        file_bytes = io.BytesIO()
        ftp.retrbinary(f'RETR {last_file}', file_bytes.write)
        ftp.quit()

        file_bytes.seek(0)

        with zipfile.ZipFile(file_bytes) as z:
            for name in z.namelist():
                if name.endswith('.xml'):
                    with z.open(name) as f:
                        try:
                            content = f.read().decode('utf-8', errors='ignore')
                            if keyword in content.lower():
                                tree = ET.ElementTree(ET.fromstring(content))
                                root = tree.getroot()

                                def find_text(tag):
                                    for elem in root.iter():
                                        if elem.tag.split('}')[-1] == tag and elem.text:
                                            return elem.text.strip()
                                    return ''

                                tender = {
                                    'file': name,
                                    'number': find_text('purchaseNumber') or find_text('notificationNumber'),
                                    'name': find_text('purchaseObjectInfo') or find_text('subject'),
                                    'price': find_text('maxPrice') or find_text('price'),
                                    'customer': find_text('fullName') or find_text('shortName'),
                                    'deadline': find_text('biddingDeadLine') or find_text('submissionCloseDateTime'),
                                    'publish_date': find_text('publishDTInEIS') or find_text('createDT'),
                                    'url': f"https://zakupki.gov.ru/epz/order/notice/ea44/view/common-info.html?regNumber={find_text('purchaseNumber')}"
                                }
                                results.append(tender)
                        except Exception:
                            continue

        return jsonify({
            'status': 'ok',
            'keyword': keyword,
            'region': region,
            'source_file': last_file,
            'count': len(results),
            'results': results[:10]
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/b2b/login', methods=['POST'])
def b2b_login():
    try:
        import requests as req
        data = request.get_json()
        login = data.get('login')
        password = data.get('password')
        resp = req.post(
            'https://www.b2b-center.ru/integration/xml/User.Login/',
            data={'login': login, 'password': password},
            timeout=15
        )
        import re
        match = re.search(r'<access_token>([^<]+)</access_token>', resp.text)
        if match:
            return jsonify({'access_token': match.group(1)})
        err = re.search(r'<message>([^<]+)</message>', resp.text)
        return jsonify({'error': err.group(1) if err else 'Ошибка авторизации'}), 401
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/b2b/request', methods=['GET'])
def b2b_request():
    try:
        import requests as req
        method = request.args.get('method')
        token = request.args.get('access_token')
        if not method or not token:
            return jsonify({'error': 'method и access_token обязательны'}), 400
        params = {k: v for k, v in request.args.items()}
        url = f'https://www.b2b-center.ru/integration/json/{method}/'
        resp = req.get(url, params=params, timeout=15)
        return jsonify(resp.json())
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)
